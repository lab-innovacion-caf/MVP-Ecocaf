import logging
import os
import html
import io
from io import BytesIO
from pypdf import PdfReader
from azure.core.credentials import AzureKeyCredential
from azure.storage.blob import BlobServiceClient
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.ai.documentintelligence import DocumentIntelligenceClient
from openai import AzureOpenAI
import json
import pandas as pd
from datetime import datetime
from azure.cosmos import CosmosClient, exceptions
import unicodedata
from docx import Document

removeall = False
remove = False
localpdfparser = False
storageaccount = os.environ["BLOB_STORAGE_ACCOUNT_NAME"]
credential = os.environ["BLOB_STORAGE_KEY"]
formrecognizerservice = os.environ["DOCUMENT_INTELLIGENCE_NAME"]
credentialformrecognizer = os.environ["DOCUMENT_INTELLIGENCE_KEY"]
source_container_name = os.environ["BLOB_SOURCE_CONTAINER_NAME"]
target_container_name = os.environ["BLOB_TARGET_CONTAINER_NAME"]
verbose = False
AZURE_OPENAI_SERVICE = os.environ["OPEN_AI_NAME"]
OPENAI_API_KEY = os.environ["OPEN_AI_KEY"]

endpoint = os.environ["COSMOS_DB_URI"]
key = os.environ["COSMOS_DB_KEY"]
database_name = os.environ["COSMOS_DB_NAME"]
container_name = os.environ["COSMOS_DB_CONTAINER_NAME"]


client = CosmosClient(endpoint, key)
try:
    # Probar la conexión a la base de datos
    database = client.get_database_client(database_name)
    logging.info('Conexión exitosa')
except exceptions.CosmosHttpResponseError as e:
    logging.error(f'Error al conectar: {e.message}')
except Exception as e:
    logging.error(f'Error al conectar: {str(e)}')


blob_service = BlobServiceClient(
    account_url=f"https://{storageaccount}.blob.core.windows.net",
    credential=credential
)
source_container_client = blob_service.get_container_client(source_container_name)



# Dynamicdictionary = {
#     "CNC": {"outputname":"CNC.csv","variables":"ACTA N°, FECHA, CLIENTE, GARANTE, ORGANISMO EJECUTOR, RIESGO INSTITUCIONAL, NOMBRE DE LA OPERACIÓN, MODALIDAD ED LA OPERACIÓN, FUENTE DE FONDOS, MONTO DE LA OPERACIÓN, PAÍS, CATEGORIA AMBIENTAL DE LA OPERACION, LINEA YSUBLINEA ESTRATÉGICA DE ACCION CAF, NIVEL DE ESTUDIOS, ESTADO DE CONTRATACION, UNIDAD DE NEGOCIO, INSTANCIA APROBATORIA, EJECUTIVO PROPONENTE DE LA OPERACIÓN, DESCRIPCIÓN DEL PROGRAMA PROYECTO, ANTECEDENTES DE LA OPERACIÓN, DESTINO DE LOS RECURSOS, PROCEDENCIA DE LOS BIENES Y SERVICIOS, Costo total (USD), Capital (USD), Préstamos (USD), Financiamiento total CAF (USD), Financiamiento Directo (USD), Cofinanciamientos/ FINANCIAMIENTO CATALITICO (USD), Otras fuentes de financiamiento (capital y/o préstamos), Promotores, Fecha estimada aprobacion,  Fecha estimada firma contrato, Cronograma tentativo de desembolsos, Nombre jurídico, Sector institucional, Sector económico del cliente, Sector económico de la operación, Rating crediticio del prestatario, Líneas Estratégicas de acción CAF, Persona contacto y cargo del cliente / organismo ejecutor, Exposición actual cliente, Exposición actual grupo economico, Exposición actual sector economico, Exposición actual pais, Exposición en Inversiones de Capital, Exposición Préstamos subordinados y otros instrumentos de cuasi capital, GRUPO ECONOMICO, Accionistas, Principales cifras de los E F (cliente), TEMAS CRÍTICOS, JUSTIFICACIÓN Y BENEFICIOS ESPERADOS DE LA OPERACIÓN, Riesgo ambiental y social, Categoría Ambiental, EQUIPO DE TRABAJO PROPUESTO, RECOMENDACIONES, Vicepresidente, Director, Director-Representante, Ejecutivos Responsable, Asistentes, Expositores, Secretario, Aspectos tratados, Recomendaciones de los miembros del Comité, Decisión, Fecha de vencimiento de al autorización"},
#     "CPI": {"outputname":"CPI.csv","variables":"ACTA N°, FECHA, FECHA DE ELABORACIÓN, FECHA DE PRESENTACIÓN DE LA OPERACION ANTE EL CNC, UNIDAD DE NEGOCIO, EJECUTIVO RESPONSABLE, EQUIPO DE TRABAJO, CLIENTE, ORGANISMO EJECUTOR, NOMBRE DE LA OPERACION, MODALIDAD DE LA OPERACION, PAIS, ACCION A SOMETER, INSTANCIA APROBATORIA, NOMBRE JURIDICO, PAIS, SECTOR INSTITUCIONAL, RIESGO INSTITUCIONAL, SECTOR ECONOMICO, EXPOSICION ACTUAL CLIENTE, NOMBRE EJECUTOR, EXPOSICION ACTUAL PAIS, OPINION DE TERCEROS, CALIFICACION DE RIESGO, FECHA DE CALIFICACION, EJECUTOR RESPONSABLE, REVISORES EXTERNOS, ASIGNADA POR EL CRC, ASPECTOS CRITICOS DE LA RELACION CON EL CLIENTE, GRUPO ECONOMICO, EXPOSICION ACTUAL CON EL GRUPO ECONOMICO, ASPECTOS CRITICOS DE LA RELACION CON EL GRUPO ECONOMICO, OPERACIONES EXISTENTES CON EL CLIENTE, TIPO DE OPERACION, MONTO, DESTINO DE LOS RECURSOS, PLAZO, TASA DE INTERES, PERIODO DE DESEMBOLSO, PERIODO DE GRACIA, AMORTIZACION, GARANTE, CONDICIONES ESPECIALES, COMISIONES, RECONOCIMIENTO DE INVERSIONES Y GASTOS, DOCUMENTACION LEGAL, EXCEPCIONES A LOS LINEAMIENTOS, OPINION DE LA DDCR, ASISTENTES, EXPOSITORES, SECRETARIO, ASPECTOS TRATADOS, RECOMENDACION DE LOS MIEMBROS DEL CPI, FECHA DE PROXIMA REVISION"},
#     "CCI" : {"outputname":"PAIS, ORGANISMOS, EJECUTOR/CLIENTE, SECTOR_INSTITUCIONAL, RIESGO_INSTITUCIONAL, NOMBRE_OPERACION, MONTO, RESUMEN_DESCRIPCION"},
#     "RRAOD" : {"outputname":"RRAOD.csv","variables":"Proyecto, Fecha de elaboración, Fecha última revisión, Equipo de trabajo, ACCIÓN A CONSIDERAR , INSTANCIA APROBATORIA DE LA REVISIÓN, Monto total aprobado, Número de operaciones en desembolso, Monto total desembolsado, EJECUCIÓN A LA FECHA DE REVISIÓN, PERFIL DE RIESGO, ANÁLISIS DE LAS OPERACIONES, CFA ASOCIADOS, Ejecutivo de Negocios, Director de Negocios, Vicepresidente de Negocios, Opinión de la VPR, APROBACIÓN DE LA REVISIÓN"},
#     "EIA" : {"outputname":"EIA.csv","variables":"PROYECTO, ETAPA ED PLANIFICACIÓN, ETAPA DE CONSTRUCCIÓN, ETAPA DE OPERACIÓN Y MANTENIMIENTO, ETAPA DE CIERRE CONSTRUCTIVO"},
#     "VPR" : {"outputname":"VPR.csv","variables":"Proyecto, País, Sector, Prestatario, Ejecutor, Fecha de Aprobación, Firma de Contrato, Inicio desembolsos, Fin desembolsos, Monto total, Monto del préstamo, Monto desembolsado, Avance Físico, Avance financiero, Equipo responsable, Numero de revision anual, Fecha de la revision anual, Descripción del proyecto, Resumen revisión previa, Area de control técnica, area de control institucional, area de control ejecución físico financiero, area de control de contratación, area de control ambiental, area de control social, area de control exposición car, area de control sostenibilidad, Indice de riesgo de ejecución"},
#     "IAS" : {"outputname":"IAS.csv","variables": "PAIS, ORGANISMOS, EJECUTOR/CLIENTE, NOMBRE_OPERACION,COMPONENTES_PROGRAMA,FINANCIAMIENTO_VERDE"},
#     "UNRECOGNIZED" : {"outputname":"UNRECOGNIZED.csv","variables":"TEXTO LIQUIDO"}
# }


def normalize_text(text):
    if text is None:
        return ""
    # Remover acentos y convertir a minúsculas
    text = unicodedata.normalize('NFKD', text).encode('ASCII', 'ignore').decode('ASCII')
    return text.lower()


def items_cosmos(pais=None, proyecto=None, tipoProyecto=None, fechaInicio=None, criterio=None, operacion=None, financiado=None, fechaFin=None, fase=None, actividad=None, categoriaP=None, categoria=None, subCategoria=None, indicadores=None, subIndicadores=None, criterios=None):
    try:
        client = CosmosClient(endpoint, key)
        database = client.get_database_client(database_name)
        container = database.get_container_client(container_name)
        excluded_fields = ["_rid", "_self", "_etag", "_attachments", "_ts"]

        # Consulta con ORDER BY descendente
        query = "SELECT * FROM c ORDER BY c.createdAt DESC"
        data = list(container.query_items(
            query=query,
            enable_cross_partition_query=True
        ))

        items = []
        for item in data:
            # Filtrar propiedades no deseadas de cada item
            filtered_item = {key: value for key, value in item.items() if key not in excluded_fields}

            # Agregar solo los datos específicos de la fase encontrada
            if "evaluacion" in filtered_item:
                items.append(filtered_item["evaluacion"])
            elif "originacion" in filtered_item:
                items.append(filtered_item["originacion"])

        # Filtrar los items después de haberlos recogido
        filtered_items = []
        for item in items:
            # Aplicamos el filtro fase específicamente para "originacion"
            if fase and fase not in item.get("FASE", ""):
                continue
            if pais and normalize_text(item.get("PAÍS", "")) != normalize_text(pais):
                continue
            if proyecto and normalize_text(proyecto) not in normalize_text(item.get("id", "")):
                continue
            if tipoProyecto and normalize_text(tipoProyecto) not in normalize_text(item.get("TIPO DE PROYECTO", "")):
                continue
            # Definir los formatos
            from datetime import datetime, timedelta

            # Para las fechas sin hora (fechaInicio, fechaFin)
            fecha_format = "%d/%m/%Y"
            # Para la fecha con hora (fecha_creacion)
            fecha_hora_format = "%d/%m/%Y %H:%M"
            logging.info("Convertir Fechas a datetime")

            # Convertir las fechas de inicio y fin
            if fechaInicio:
                fechaInicio2 = datetime.strptime(fechaInicio, fecha_format)
            if fechaFin:
                fechaFin2 = datetime.strptime(fechaFin, fecha_format)

            # Si fechaInicio y fechaFin son iguales, ajusta para incluir todo el día
            if fechaInicio and fechaFin and fechaInicio == fechaFin:
                fechaInicio2 = fechaInicio2.replace(hour=0, minute=0)
                fechaFin2 = fechaFin2.replace(hour=23, minute=59)

            logging.info(fechaInicio)

            # Convertir la fecha de creación (que incluye hora)
            fecha_creacion = item.get("FECHA CREACION", "01/01/1900 00:00")

            if fecha_creacion != "01/01/1900 00:00" and (fechaInicio or fechaFin):
                fecha_creacion = datetime.strptime(fecha_creacion, fecha_hora_format)

                # Filtros
                if fechaInicio and fechaFin:
                    logging.info("Fecha inicio y fin")
                    logging.info(fechaFin2)
                    logging.info(fechaInicio2)
                    if not (fechaInicio2 <= fecha_creacion <= fechaFin2):
                        continue

                if fechaInicio and fecha_creacion < fechaInicio2:
                    logging.info("Fecha creacion menor a fecha inicio")
                    continue

            if criterio and criterio not in item.get("ELEGIBLE / NO ELEGIBLE", ""):
                continue
            if operacion and operacion not in item.get("ES SOBERANO / NO SOBERANO", ""):
                continue
            if financiado and normalize_text(financiado) not in normalize_text(item.get("FINANCIADO POR?", "")):
                continue
            if actividad:
                actividades_proyecto = item.get("ACTIVIDADES ELEGIBLES QUE APLICAN AL PROYECTO", "")
                actividades_list = [item2.strip() for item2 in actividades_proyecto.split('|')]
                logging.info(f'Actividades: {actividades_list}')
                # Verifica si la actividad.strip() está en la lista de actividades
                if actividad.strip().lower() not in [act.lower() for act in actividades_list]:
                    continue

            if categoriaP:
                categoriaP_proyecto = item.get("CATEGORÍAS PRINCIPALES DE FINANCIAMIENTO VERDE EN LAS QUE CLASIFICA", "")
                categoriaP_list = [item2.strip() for item2 in categoriaP_proyecto.split('|')]
                logging.info(f'categoria: {categoriaP_list}')
                # Verifica si la actividad.strip() está en la lista de actividades
                if categoriaP.strip().lower() not in [act.lower() for act in categoriaP_list]:
                    continue

            if categoria:
                categoria_proyecto = item.get("CATEGORÍAS DE FINANCIAMIENTO VERDE EN LAS QUE CLASIFICA", "")
                categoria_list = [item2.strip() for item2 in categoria_proyecto.split(',')]
                logging.info(f'categoria: {categoria_list}')
                # Verifica si la actividad.strip() está en la lista de actividades
                if categoria.strip().lower() not in [act.lower() for act in categoria_list]:
                    continue

            if subCategoria:
                subCategoria_proyecto = item.get("SUBCATEGORÍAS DE FINANCIAMIENTO VERDE EN LAS QUE CLASIFICA", "")
                subCategoria_list = [item2.strip() for item2 in subCategoria_proyecto.split(',')]
                logging.info(f'subCat: {subCategoria_list}')
                # Verifica si la actividad.strip() está en la lista de actividades
                if subCategoria.strip().lower() not in [act.lower() for act in subCategoria_list]:
                    continue

            if indicadores:
                indicadores_proyecto = item.get("INDICADORES", "")
                indicadores_list = [item2.strip() for item2 in indicadores_proyecto.split('|')]
                logging.info(f'subCat: {indicadores_list}')
                # Verifica si la actividad.strip() está en la lista de actividades
                if indicadores.strip().lower() not in [act.lower() for act in indicadores_list]:
                    continue

            if subIndicadores:
                subindicadores_proyecto = item.get("SUBINDICADORES", "")
                subindicadores_list = [item2.strip() for item2 in subindicadores_proyecto.split('|')]
                logging.info(f'subCat: {subindicadores_list}')
                # Verifica si la actividad.strip() está en la lista de actividades
                if subIndicadores.strip().lower() not in [act.lower() for act in subindicadores_list]:
                    continue

            if criterios:
                criterios_proyecto = item.get("CRITERIOS DE ELEGIBILIDAD", "")
                criterios_list = [item2.strip() for item2 in criterios_proyecto.split('|')]
                logging.info(f'subCat: {criterios_list}')
                # Verifica si la actividad.strip() está en la lista de actividades
                if criterios.strip().lower() not in [act.lower() for act in criterios_list]:
                    continue

            filtered_items.append(item)

        return filtered_items
    except exceptions.CosmosHttpResponseError as e:
        logging.error(f'Error al conectar: {e.message}')
        return f'Error al conectar: {e.message}'
    except Exception as e:
        logging.error(f'Error al conectar: {str(e)}')
        return f'Error al conectar: {str(e)}'


def get_items_by_id(item_id):
    try:
        # Inicializar el cliente de Cosmos DB
        client = CosmosClient(endpoint, key)
        database = client.get_database_client(database_name)
        container = database.get_container_client(container_name)
        excluded_fields = ["_rid", "_self", "_etag", "_attachments", "_ts"]

        # Almacenar el resultado
        # Asegúrate de tener la clave de partición correcta
        item = container.read_item(item_id, partition_key=item_id)

        # Filtrar los campos excluidos
        if item:
            filtered_item = {key: value for key, value in item.items() if key not in excluded_fields}
            return filtered_item

        return None  # Retornar None si no se encuentra el ítem
    except exceptions.CosmosHttpResponseError as e:
        return f'Error al conectar: {e.message}'
    except Exception as e:
        return f'Error al conectar: {str(e)}'


def convertir_a_float(valor):
    try:
        # Eliminar texto adicional y reemplazar coma con punto decimal
        valor = valor.replace("USD", "").replace("MM", "").replace(",", "").strip()
        return float(valor)
    except ValueError:
        return 0


def limpiar_nulls(data):
    # Recorre el diccionario y reemplaza los valores None por cadenas vacías
    for key, value in data.items():
        if isinstance(value, dict):
            # Llamada recursiva para limpiar diccionarios anidados
            limpiar_nulls(value)
        elif value is None:
            data[key] = ""
    return data


def insertar_item_cosmos(data):
    try:
        # Validar que los datos contengan los campos necesarios
        if 'id' not in data:
            return "Error: El ítem debe tener un campo 'id'"

        # Limpiar los valores None en el data
        data = limpiar_nulls(data)

        client = CosmosClient(endpoint, key)
        database = client.get_database_client(database_name)
        container = database.get_container_client(container_name)

        # Obtener el documento existente (si existe) para mantener las fases anteriores
        existing_item = None
        try:
            existing_item = container.read_item(item=data['id'], partition_key=data['id'])
        except exceptions.CosmosResourceNotFoundError:
            # Crear un nuevo documento si no existe
            existing_item = {'id': data['id']}

        # Determinar la fase actual y almacenar los datos en el campo correspondiente
        fase_actual = data.get('FASE', '')

        if fase_actual == "originacion":
            existing_item['originacion'] = data
        elif fase_actual == "evaluacion":
            existing_item['evaluacion'] = data
        else:
            return "Error: Fase desconocida. Debe ser 'originacion', 'evaluacion'."

        # Insertar o actualizar el documento en Cosmos DB
        container.upsert_item(existing_item)
        return 'Item insertado o actualizado exitosamente'
    except exceptions.CosmosHttpResponseError as e:
        return f'Error al conectar: {e.message} - Código de error: {e.status_code}'
    except Exception as e:
        return f'Error inesperado: {str(e)}'


def update_item_cosmos(data):
        if 'id' not in data:
            return "Error: El ítem debe tener un campo 'id'"

        # Limpiar los valores None en el data
        data = limpiar_nulls(data)

        client = CosmosClient(endpoint, key)
        database = client.get_database_client(database_name)
        container = database.get_container_client(container_name)

        get_items_by_id = container.read_item(item=data['id'], partition_key=data['id'])
        if 'originacion' in get_items_by_id:
            get_items_by_id['originacion'] = data['originacion']

        if 'evaluacion' in get_items_by_id:
            get_items_by_id['evaluacion'] = data['evaluacion']

        if 'biodiversidad' in data:
            get_items_by_id['biodiversidad'] = data['biodiversidad']

        container.upsert_item(get_items_by_id)

def table_to_html(table):
    """
    Converts a table object to an HTML table string.

    Parameters:
    table (Table): The table object to convert.

    Returns:
    str: The HTML table string.

    """
    # Initialize the HTML table.
    table_html = "<table>"

    # Sort the cells by row and column.
    rows = [sorted([cell for cell in table.cells if cell.row_index == i], key=lambda cell: cell.column_index) for i in range(table.row_count)]

    # Convert each cell to an HTML tag and add it to the table.
    for row_cells in rows:
        table_html += "<tr>"
        for cell in row_cells:
            # Determine the tag type based on the cell kind.
            tag = "th" if (cell.kind == "columnHeader" or cell.kind == "rowHeader") else "td"

            # Add the cell's row and column spans to the tag.
            cell_spans = ""
            if (cell.column_span is not None):
                if cell.column_span > 1:
                    cell_spans += f" colSpan={cell.column_span}"

            if (cell.row_span is not None):
                if cell.row_span > 1:
                    cell_spans += f" rowSpan={cell.row_span}"

            # Convert the cell content to HTML and add it to the table.
            table_html += f"<{tag}{cell_spans}>{html.escape(cell.content)}</{tag}>"

        # Close the row tag.
        table_html += "</tr>"

    # Close the table tag and return the HTML string.
    table_html += "</table>"
    return table_html


def find_dictionary(words, dynamic_dict, filename):
    for word in words:
        if word.strip("()") in dynamic_dict:
            return dynamic_dict[word.strip("()")]
        if word.strip("[]") in dynamic_dict:
            return dynamic_dict[word.strip("[]")]

    namewords = filename.split()
    for word in namewords:
        if word.strip("()") in dynamic_dict:
            return dynamic_dict[word.strip("()")]
        if word.strip("[]") in dynamic_dict:
            return dynamic_dict[word.strip("[]")]
    namewordsdownslash = filename.split("_")
    for word in namewordsdownslash:
        if word.strip("()") in dynamic_dict:
            return dynamic_dict[word.strip("()")]
        if word.strip("[]") in dynamic_dict:
            return dynamic_dict[word.strip("[]")]

    return dynamic_dict["UNRECOGNIZED"]

# def get_document_text(filename):
    offset = 0
    page_map = []
    Titles = []
    dictionary = {}

    # Create a BlobServiceClient using the account URL and credentials.
    blob_service = BlobServiceClient(
        account_url=f"https://{storageaccount}.blob.core.windows.net",
        credential=credential
    )
    logging.info(f"Processing file: {filename}")
    if localpdfparser:
        # Use local PDF parser to extract text
        blob_client = blob_service.get_blob_client(container=source_container_name, blob=filename)
        # Download the PDF file as bytes and create a BytesIO object.
        data = blob_client.download_blob().readall()
        logging.info(f"Downloaded {len(data)} bytes from {filename}.")
        reader = PdfReader(io.BytesIO(data))
        pages = reader.pages
        for page_num, p in enumerate(pages):
            page_text = p.extract_text()
            page_map.append((page_num, offset, page_text))
            offset += len(page_text)
    else:
        # Use Azure Form Recognizer to extract text
        if verbose:
            logging.info(f"Extracting text from '{filename}' using Azure Form Recognizer")
            logging.info(f"Extracting text from '{filename}' using Azure Form Recognizer")
        if os.path.splitext(filename)[1].lower() == ".pdf":
            form_recognizer_client = DocumentAnalysisClient(
                endpoint=f"https://{formrecognizerservice}.cognitiveservices.azure.com/",
                credential=AzureKeyCredential(credentialformrecognizer),
                headers={"x-ms-useragent": "azure-search-chat-demo/1.0.0"}
            )

            blob_client = blob_service.get_blob_client(container=source_container_name, blob=filename)

            poller = form_recognizer_client.begin_analyze_document("prebuilt-layout", document=blob_client.download_blob().content_as_bytes())

        if os.path.splitext(filename)[1].lower() == ".docx":
            document_intelligence_client = DocumentIntelligenceClient(
                endpoint=f"https://{formrecognizerservice}.cognitiveservices.azure.com/",
                credential=AzureKeyCredential(credentialformrecognizer)
            )
            blob_client = blob_service.get_blob_client(container=source_container_name, blob=filename)

            poller = document_intelligence_client.begin_analyze_document(
                model_id="prebuilt-layout",
                analyze_request=blob_client.download_blob().content_as_bytes(),
                content_type="application/octet-stream"
            )
        try:
            form_recognizer_results = poller.result()
            Titles = ' '.join([parregraph.content for parregraph in form_recognizer_results.paragraphs[1:6]])
            dictionary = find_dictionary(Titles.split(), Dynamicdictionary, filename)
        except:
            raise Exception("Documento no soportado!!")

        for page_num, page in enumerate(form_recognizer_results.pages):
            tables_on_page = []
            # Mark all positions of the table spans in the page
            page_offset = page.spans[0].offset
            page_length = page.spans[0].length
            table_chars = [-1] * page_length
            if os.path.splitext(filename)[1].lower() == ".pdf":
                tables_on_page = [table for table in form_recognizer_results.tables if table.bounding_regions[0].page_number == page_num + 1]
            if os.path.splitext(filename)[1].lower() == ".docx":
                for table in form_recognizer_results.tables:
                    if table.spans:
                        table_start = table.spans[0].offset
                        table_end = table_start + table.spans[0].length
                        if page_offset <= table_start < page_offset + page_length:
                            tables_on_page.append(table)
                            table_id = len(tables_on_page) - 1
                            for span in table.spans:
                                for i in range(span.length):
                                    idx = span.offset - page_offset + i
                                    if 0 <= idx < page_length:
                                        table_chars[idx] = table_id

            for table_id, table in enumerate(tables_on_page):
                for span in table.spans:
                    # Replace all table spans with "table_id" in table_chars array
                    for i in range(span.length):
                        idx = span.offset - page_offset + i
                        if idx >= 0 and idx < page_length:
                            table_chars[idx] = table_id

            # Build page text by replacing characters in table spans with table html
            page_text = ""
            added_tables = set()
            for idx, table_id in enumerate(table_chars):
                if table_id == -1:
                    page_text += form_recognizer_results.content[page_offset + idx]
                elif not table_id in added_tables:
                    page_text += table_to_html(tables_on_page[table_id])
                    added_tables.add(table_id)
            page_text += " "
            page_map.append((page_num, offset, page_text))
            offset += len(page_text)

    return page_map, dictionary


def get_document_text(filename):
    # Descargar el archivo desde Azure Blob Storage
    blob_client = blob_service.get_blob_client(container=source_container_name, blob=filename)
    data = blob_client.download_blob().readall()
    logging.info(f"Downloaded {len(data)} bytes from {filename}.")

    # Detectar el tipo de archivo basándonos en la extensión
    file_extension = os.path.splitext(filename)[1].lower()
    file_path = os.path.splitext(filename)[0]
    output_txt_filename = f"{file_path}.txt"

    if file_extension == '.pdf':
        # Usar Azure Form Recognizer para extraer texto de archivos PDF
        form_recognizer_client = DocumentAnalysisClient(
            endpoint=f"https://{formrecognizerservice}.cognitiveservices.azure.com/",
            credential=AzureKeyCredential(credentialformrecognizer)
        )

        poller = form_recognizer_client.begin_analyze_document("prebuilt-document", data)
        result = poller.result()

        # Procesar resultados del PDF
        document_text = ""
        for page in result.pages:
            document_text += f"--- Page {page.page_number} ---\n"
            for line in page.lines:
                # Accede al contenido de cada línea
                document_text += f"{line.content}\n"

    elif file_extension == '.docx':
        # Usar python-docx para extraer texto de archivos .docx
        document_text = ""
        doc = Document(BytesIO(data))  # Lee el contenido desde BytesIO
        for paragraph in doc.paragraphs:
            document_text += paragraph.text + "\n"  # Agrega el contenido de cada párrafo

    else:
        raise ValueError(f"Formato de archivo no soportado: {file_extension}")

    # Guardar el texto extraído en un archivo .txt en Azure Blob Storage
    # Convertir el texto a bytes
    txt_data = BytesIO(document_text.encode('utf-8'))
    txt_blob_client = blob_service.get_blob_client(container=source_container_name, blob=output_txt_filename)
    logging.info("Uploading txt")
    txt_blob_client.upload_blob(txt_data, overwrite=True)
    return document_text


def updatetostorage(df, outputname):
    blob_service = BlobServiceClient(
        account_url=f"https://{storageaccount}.blob.core.windows.net",
        credential=credential
    )

    # Get target container Client
    target_container_client = blob_service.get_container_client(target_container_name)
    if not target_container_client.exists():
        target_container_client.create_container()

    # Get the blob client for the PDF file.
    blob_client = target_container_client.get_blob_client(outputname)
    if blob_client.exists():
        # Descargar el blob existente a un DataFrame
        blob_data = blob_client.download_blob().content_as_bytes()
        existing_df = pd.read_csv(BytesIO(blob_data))

        # Añadir el nuevo DataFrame al DataFrame existente
        combined_df = pd.concat([existing_df, df])
    else:
        # Si el blob no existe, usar el nuevo DataFrame
        combined_df = df

    # Guardar el DataFrame combinado a un CSV en una cadena
    csv_data = combined_df.to_csv(index=False)

    # Subir el CSV actualizado al blob
    blob_client.upload_blob(csv_data, overwrite=True)

# def processdocument(filename):
    try:
        page_map, dictionary = get_document_text(filename)
    except Exception as e:
        logging.error(f"Error al obtener el texto del documento: {e}")
        return None, None, None

    outputname = f"OUTPUT/{dictionary['outputname']}"
    variablesarray = dictionary.get("variables", [])

    if not variablesarray:
        logging.info("No se encontraron variables en el diccionario.")
        return None, None, None

    columnas = [var["nombre"] for var in variablesarray]
    Variables = " | ".join(columnas)
    logging.info(f"Document type: {dictionary['outputname']}")

    Allcontext = ''.join([f"{page[2]}" for page in page_map])

    logging.info("Connecting to OpenAI")
    try:
        client = AzureOpenAI(
            api_key=OPENAI_API_KEY,
            azure_endpoint=f"https://{AZURE_OPENAI_SERVICE}.openai.azure.com",
            api_version="2023-03-15-preview"
        )

        chat_prompt = [
            {"role": "system", "content": "Eres un asistente que ayuda a extraer campos de documentos..."},
            {"role": "user", "content": f"Este es el contexto: {Allcontext}, Estos son los campos: {Variables}"}
        ]

        logging.info("Querying OpenAI")
        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=chat_prompt,
            temperature=0.5,
            frequency_penalty=0,
            presence_penalty=0
        )
        logging.info("OpenAI completed")

        answer = completion.choices[0].message.content
        logging.info("PARSING RESULT")

        json_string = answer.strip('```json\n')
        logging.info(f"RESULT TO PARSE: {json_string}")
        result = json.loads(json_string)

        df = pd.DataFrame([result])  # Si necesitas trabajar con un DataFrame
        logging.info("UPLOADING RESULT")
        # updatetostorage(df, outputname)  # Descomentar cuando estés listo para subir

        return result, dictionary['outputname'].split(".")[0], outputname

    except Exception as e:
        logging.error(f"Ocurrió un error durante el procesamiento con OpenAI: {e}")
        return None, None, None


def processdocument(Allcontext, document_type):
    try:
        
        blob = source_container_client.get_blob_client(f"Parametrizaciones/Dict.json").download_blob().readall()
        json_data = json.loads(blob)
        Dynamicdictionary = json_data

        logging.info("Connecting to OpenAI")
        client = AzureOpenAI(
            api_key=OPENAI_API_KEY,
            azure_endpoint=f"https://{AZURE_OPENAI_SERVICE}.openai.azure.com",
            api_version="2023-03-15-preview"
        )

        variablesarray = Dynamicdictionary[document_type].get("variables", [])

        chat_prompt = [
            {"role": "system", "content": "Eres un asistente especializado en extraer información estructurada de documentos y presentarla en un formato JSON. Los campos por extraer están en un array, cada uno con un nombre de campo y una descripción. Para los campos que incluyen montos, convierte el monto a un valor numérico sin unidades, si el monto está en formato como 'USD 60 M', convierte el valor a su representación numérica en base 10 por ejemplo USD 60 M a 60000000. Al brindar la información extraída ten en cuenta el tipo de dato que se indica en la descripción de cada campo."},
            {"role": "user", "content": f"Este es el contexto: {Allcontext}. Por favor, extrae los siguientes campos: {variablesarray}"}
        ]

        logging.info("Querying OpenAI")
        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=chat_prompt,
            temperature=0.2,
            frequency_penalty=0,
            presence_penalty=0
        )
        logging.info("OpenAI completed")

        answer = completion.choices[0].message.content
        logging.info("PARSING RESULT")

        # Limpiar la respuesta para asegurarse de que sea JSON válido
        json_string = answer.strip()

        # Eliminar el prefijo ```json y el sufijo ``` si están presentes
        if json_string.startswith("```json"):
            json_string = json_string[7:].strip()  # Eliminar "```json\n"
        if json_string.endswith("```"):
            json_string = json_string[:-3].strip()  # Eliminar "```"

        logging.info(f"RESULT TO PARSE: {json_string}")

        try:
            result = json.loads(json_string)
        except json.JSONDecodeError as json_error:
            logging.info(f"Error al analizar JSON: {json_error}")
            return None

        # Convertir el resultado a un DataFrame si es necesario
        df = pd.DataFrame([result])
        logging.info(f"RESULT DATAFRAME: {df}")

        return result
    except Exception as e:
        logging.error(f"Ocurrió un error durante el procesamiento con OpenAI: {e}")
        raise ValueError(f"[processdocument] - Error: {str(e)}") 


# def extractdocument(file_path):
    # Obtener la referencia del archivo procesado en el blob storage
    proceseddocuments = source_container_client.get_blob_client("LOGS/PROCESSED_DOCUMENTS.csv")

    # Cargar el archivo de documentos procesados si existe
    if proceseddocuments.exists():
        blob_data = proceseddocuments.download_blob().content_as_bytes()
        proceseddocuments_df = pd.read_csv(BytesIO(blob_data))
    else:
        proceseddocuments_df = pd.DataFrame(columns=["nombre", "fecha"])

    documentos_errores = []
    documentos_procesados = 0

    # Obtener la extensión del archivo
    file_extension = os.path.splitext(file_path)[1].lower()
    logging.info(f"File extension: {file_extension}")

    # Procesar solo archivos .pdf o .docx
    if file_extension in [".pdf", ".docx"]:
        try:
            logging.info(f"Processing.... {file_path}")

            # Verificar si el documento ya ha sido procesado
            has_been_processed = os.path.basename(file_path) in proceseddocuments_df["nombre"].values
            logging.info(f"Processed? {has_been_processed}")

            if not has_been_processed:
                documentos_procesados += 1
                # Procesar el documento y obtener los resultados
                # Aquí se llamaría a tu función que procesa el documento
                document_text = get_document_text(file_path)
                result, doc_type, outputname = processdocument(document_text)
                return {"variables": result, "tipo": doc_type, "outputname": outputname}
            else:
                logging.info(f"{file_path} has already been processed.")
                return None  # No procesar de nuevo si ya ha sido procesado
        except Exception as e:
            logging.error("Error during processing")
            logging.error(e)
            documentos_errores.append(file_path)
            return None  # Error al procesar el documento

    else:
        logging.info(f"Unsupported file format: {file_extension}")
        return None


def extractdocument(filename, document_type):
    # Descargar el archivo y obtener su texto
    document_text = get_document_text(filename)
    # Procesar el texto del documento y extraer los campos requeridos
    fields = processdocument(document_text, document_type)

    # Retornar los campos extraídos
    logging.info(f"Extraídos: {fields}")
    return fields


def savedocumentdata(filedata):
    logging.info("StartProcessing")
    result = filedata.get("variables")
    df = pd.DataFrame([result])
    logging.info("UPLOADING RESULT")
    updatetostorage(df, filedata.get("outputname"))
    proceseddocuments = source_container_client.get_blob_client("LOGS/PROCESED_DOCUMENTS.csv")

    if proceseddocuments.exists():
        # Descargar el blob existente a un DataFrame
        blob_data = proceseddocuments.download_blob().content_as_bytes()
        proceseddocuments_df = pd.read_csv(BytesIO(blob_data))

    else:
        # Si el blob no existe, usar el nuevo DataFrame
        proceseddocuments_df = pd.DataFrame(columns=["nombre", "fecha"])

    new_row = pd.DataFrame([{"nombre": filedata.get("filename"), "fecha": datetime.now().strftime("%Y-%m-%d")}])
    proceseddocuments_df = pd.concat([proceseddocuments_df, new_row], ignore_index=True)
    csv_data = proceseddocuments_df.to_csv(index=False)
    proceseddocuments.upload_blob(csv_data, overwrite=True)
    return {"result": "ok"}


def getdata(filename):
    proceseddocuments = source_container_client.get_blob_client(f"OUTPUT/{filename}")

    if proceseddocuments.exists():
        # Descargar el blob existente a un DataFrame
        blob_data = proceseddocuments.download_blob().content_as_bytes()
        proceseddocuments_df = pd.read_csv(BytesIO(blob_data))

        return proceseddocuments_df.to_json(orient="records", indent=4)

    else:
        # Si el blob no existe, usar el nuevo DataFrame
        return {"result": "Error"}


def export_data_to_excel(anio_filtrar):

    database = client.get_database_client(database_name)

    # Construir la consulta para filtrar por año
    query = f"SELECT * FROM c WHERE c.Fecha LIKE '%{anio_filtrar}'"

    # Obtener los documentos que coinciden con la consulta
    containerclient = database.get_container_client("Proyectos")
    items = list(containerclient.query_items(query=query, enable_cross_partition_query=True))

    # Convertir los resultados a un array de diccionarios
    resultados_array = [dict(item) for item in items]

    # Convertir el array de diccionarios en un DataFrame dinámico
    df = pd.DataFrame(resultados_array)

    # Eliminar columnas no deseadas
    columns_to_remove = ['_rid', '_self', '_etag', '_attachments', '_ts']
    df = df.drop(columns=columns_to_remove, errors='ignore')

    # Exportar el DataFrame a un archivo Excel (.xlsx) con el año incluido en el nombre
    xlsx_filename = f'datos_proyectos_{anio_filtrar}.xlsx'
    df.to_excel(xlsx_filename, index=False)

    return xlsx_filename
